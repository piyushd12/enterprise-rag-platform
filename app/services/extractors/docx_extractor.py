import io
import logging
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)


class DOCXExtractor:
    def extract(self, file_bytes: bytes) -> dict:
        doc = DocxDocument(io.BytesIO(file_bytes))
        parts = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    parts.append(row_text)

        full_text = "\n\n".join(parts)
        logger.info(f"DOCX extracted: {len(full_text)} characters, {len(parts)} text blocks")

        return {
            "text": full_text,
            "page_count": 1,
            "pages": [{"page_num": 1, "text": full_text}],
        }