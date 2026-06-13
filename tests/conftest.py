import pytest
import pytest_asyncio
from httpx import ASGITransport, AsyncClient
from sqlalchemy.ext.asyncio import AsyncSession, async_sessionmaker, create_async_engine
import fitz

from app.core.database import get_db
from app.main import app
from app.models.base import Base

TEST_DB_URL = "sqlite+aiosqlite:///./test.db"

test_engine = create_async_engine(TEST_DB_URL, echo=False)
TestSessionFactory = async_sessionmaker(
    bind=test_engine, class_=AsyncSession, expire_on_commit=False
)


@pytest_asyncio.fixture(scope="session", autouse=True)
async def setup_test_database():
    async with test_engine.begin() as conn:
        await conn.run_sync(Base.metadata.create_all)
    yield
    async with test_engine.begin() as conn:
        await conn.run_sync(Base.metadata.drop_all)
    await test_engine.dispose()


@pytest_asyncio.fixture
async def db():
    async with TestSessionFactory() as session:
        yield session
        await session.rollback()


@pytest_asyncio.fixture
async def client(db: AsyncSession):

    async def override_db():
        yield db

    app.dependency_overrides[get_db] = override_db

    async with AsyncClient(
        transport=ASGITransport(app=app), base_url="http://test"
    ) as c:
        yield c

    app.dependency_overrides.clear()


@pytest_asyncio.fixture
async def registered_user(client: AsyncClient) -> dict:
    response = await client.post("/auth/register", json={
        "email": "testuser@example.com",
        "password": "testpassword123",
        "full_name": "Test User",
    })
    assert response.status_code == 201
    return response.json()


@pytest_asyncio.fixture
async def auth_token(client: AsyncClient, registered_user: dict) -> str:
    response = await client.post("/auth/login", json={
        "email": "testuser@example.com",
        "password": "testpassword123",
    })
    return response.json()["access_token"]


@pytest_asyncio.fixture
def auth_headers(auth_token: str) -> dict:
    return {"Authorization": f"Bearer {auth_token}"}


@pytest.fixture
def sample_pdf_bytes() -> bytes:
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((50, 100), "The Eiffel Tower is 330 meters tall.")
    page.insert_text((50, 150), "It is located in Paris, France.")
    return doc.tobytes()


@pytest.fixture
def sample_docx_bytes() -> bytes:
    from docx import Document as DocxDocument
    import io
    doc = DocxDocument()
    doc.add_paragraph("Artificial intelligence is transforming the world.")
    doc.add_paragraph("Machine learning is a subset of AI.")
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()